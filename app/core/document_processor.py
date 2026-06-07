"""Document parsing pipeline for PDF, DOCX, and plain-text files."""

import io
import logging
from dataclasses import dataclass

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.config import settings
from app.utils.text_utils import clean_text, split_into_chunks

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class ParsedDocument:
    """Immutable result of parsing a document file.

    Attributes:
        filename: Original uploaded filename.
        raw_text: Cleaned full text extracted from the document.
        chunks: List of overlapping text chunks ready for embedding.
        page_count: Number of pages (PDF) or 1 for DOCX/TXT.
        char_count: Total character count of raw_text.
    """

    filename: str
    raw_text: str
    chunks: list[str]
    page_count: int
    char_count: int


def _parse_pdf(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from a PDF byte stream.

    Args:
        file_bytes: Raw bytes of the PDF file.

    Returns:
        A tuple of (extracted_text, page_count).
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages: list[str] = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n\n".join(pages), len(pages)


def _parse_docx(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from a DOCX byte stream.

    Args:
        file_bytes: Raw bytes of the DOCX file.

    Returns:
        A tuple of (extracted_text, page_count) where page_count is always 1.
    """
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs), 1


def _parse_txt(file_bytes: bytes) -> tuple[str, int]:
    """Decode a plain-text byte stream.

    Attempts UTF-8 first, falls back to latin-1 on decode errors.

    Args:
        file_bytes: Raw bytes of the text file.

    Returns:
        A tuple of (decoded_text, page_count) where page_count is always 1.
    """
    try:
        return file_bytes.decode("utf-8"), 1
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1"), 1


def parse_document(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Parse a document file into cleaned text and embedding-ready chunks.

    Routes to the appropriate parser based on the file extension, then applies
    text cleaning and chunking using the application settings.

    Args:
        file_bytes: Raw bytes of the uploaded file.
        filename: Original filename, used to determine the parser and stored
            in the returned dataclass.

    Returns:
        A ParsedDocument containing the raw text, chunks, and metadata.

    Raises:
        ValueError: If the file extension is not supported or the extracted
            text is empty after cleaning.
    """
    suffix = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    _parsers = {"pdf": _parse_pdf, "docx": _parse_docx, "txt": _parse_txt}

    if suffix not in _parsers:
        raise ValueError(f"Unsupported file type: .{suffix}. Accepted: pdf, docx, txt.")

    raw, page_count = _parsers[suffix](file_bytes)
    logger.debug("Parsed %s: %d chars, %d pages", filename, len(raw), page_count)

    cleaned = clean_text(raw)
    if not cleaned:
        raise ValueError(f"No text could be extracted from '{filename}'.")

    chunks = split_into_chunks(cleaned, settings.max_chunk_size, settings.chunk_overlap)
    logger.info("Document '%s' → %d chunks", filename, len(chunks))

    return ParsedDocument(
        filename=filename,
        raw_text=cleaned,
        chunks=chunks,
        page_count=page_count,
        char_count=len(cleaned),
    )
